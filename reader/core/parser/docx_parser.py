"""DOCX (Word 2007+) book import.

Word documents carry real paragraph styles, so when an author used Heading 1
the chapter boundaries are stated rather than guessed. Documents without any
heading styles fall back to the same text heuristics the TXT importer uses.

python-docx is imported lazily inside the functions that need it so the app
still starts when the dependency is missing.
"""

import re

from core.parser.language import detect_language
from core.parser.sections import (
    EXPLICIT_MARKER_THRESHOLD,
    build_chapters,
    is_explicit_section,
)


class DocxImportError(RuntimeError):
    """A .docx file could not be read, with a message meant for the user."""


# Word's built-in heading styles are "Heading 1" through "Heading 9". "Title"
# is the document title style and also marks a boundary.
_HEADING_STYLE_PREFIX = 'Heading'
_TITLE_STYLE = 'Title'


def _is_heading_style(style_name: str) -> bool:
    name = (style_name or '').strip()
    return name.startswith(_HEADING_STYLE_PREFIX) or name == _TITLE_STYLE


def _walk_container(container, doc, units, in_table):
    """Recursively collect (text, is_heading) pairs from a body-like element.

    Descends into table cells (picking up nested tables too) and into
    content-control (w:sdt) wrappers via their w:sdtContent, since both are
    otherwise silently skipped by a body-level-only walk. in_table disables
    heading detection for anything nested inside a table cell: a heading-
    styled cell is a table header, not a chapter boundary.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in container.iterchildren():
        if not isinstance(child.tag, str):
            # lxml yields comments/processing instructions with a callable
            # .tag rather than a string; they carry no book text.
            continue
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            paragraph = Paragraph(child, doc)
            if in_table and not paragraph.text.strip():
                # Blank paragraphs inside a table cell are layout padding, not
                # a narration break, so unlike top-level blanks they are
                # dropped rather than kept as blank units.
                continue
            is_heading = (not in_table) and _is_heading_style(paragraph.style.name)
            units.append((paragraph.text, is_heading))
        elif tag == 'tbl':
            # Table text is body content, gathered cell by cell in row order.
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    _walk_container(cell._tc, doc, units, True)
        elif tag == 'sdt':
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                _walk_container(sdt_content, doc, units, in_table)
        # Anything else (sectPr, bookmarks) carries no book text.


def extract_units(doc):
    """Return ordered (text, is_heading) pairs for a python-docx Document.

    Walks the document body rather than doc.paragraphs, because
    doc.paragraphs silently omits text inside tables and content controls.
    Top-level blank paragraphs are kept as blank units (they mark a narration
    break), but blank paragraphs inside a table cell are filtered out, since
    table layout padding is not a narration break.
    """
    units = []
    _walk_container(doc.element.body, doc, units, False)
    return units


def has_chapter_heading_styles(doc) -> bool:
    """True when the document uses a real Heading style for structure.

    Title alone does not count: a Title-styled title page is common in
    documents whose chapter headings are plain text, and treating it as
    declared structure would collapse the whole book into one chapter.
    """
    for paragraph in doc.paragraphs:
        name = (paragraph.style.name or '').strip()
        if name.startswith(_HEADING_STYLE_PREFIX):
            return True
    return False


# Legacy .doc files are OLE containers and start with this signature. Renaming
# one to .docx is the most likely user mistake, and python-docx reports it as
# "Package not found", which reads like a missing file.
_OLE_MAGIC = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'


def _is_legacy_doc(file_path) -> bool:
    try:
        with open(file_path, 'rb') as handle:
            return handle.read(len(_OLE_MAGIC)) == _OLE_MAGIC
    except OSError:
        return False


def _open_document(file_path):
    """Open a .docx, translating library errors into user-facing messages."""
    if _is_legacy_doc(file_path):
        raise DocxImportError(
            'This is a legacy .doc file. Open it in Word and save as .docx.'
        )
    try:
        from docx import Document
    except ImportError as exc:
        raise DocxImportError(
            'DOCX support needs python-docx. Run: pip install python-docx'
        ) from exc
    try:
        return Document(file_path)
    except DocxImportError:
        raise
    except Exception as exc:
        raise DocxImportError(
            'Could not read the .docx file. It may be corrupt or password protected.'
        ) from exc


# python-docx stamps this as the author of any document it creates, so a file
# generated by a script would otherwise import as written by "python-docx".
_GENERATOR_AUTHORS = {'python-docx'}


def _metadata(doc, units):
    """Title and author from core properties, falling back to the text."""
    props = doc.core_properties
    title = (getattr(props, 'title', None) or '').strip()
    author = (getattr(props, 'author', None) or '').strip()
    if author in _GENERATOR_AUTHORS:
        author = ''

    if not title:
        title = 'Unknown Title'
        for text, _ in units[:20]:
            stripped = (text or '').strip()
            if stripped and len(stripped) < 120:
                title = stripped
                break

    if not author:
        author = 'Unknown Author'
        head = '\n'.join(text for text, _ in units[:20])[:500]
        match = re.search(r'\bby[ \t]+([A-Z][a-z]+(?:[ \t]+[A-Z][a-z]+)*)', head)
        if match:
            author = match.group(1)

    return title, author


def parse(file_path):
    doc = _open_document(file_path)
    units = extract_units(doc)

    raw = '\n'.join(text for text, _ in units)
    if not raw.strip():
        raise DocxImportError('No readable text found in the .docx file.')

    title, author = _metadata(doc, units)

    # Either/or: real heading styles fully replace text-based detection, so an
    # author's declared structure is never second-guessed by a regex. Title
    # alone does not arm this: it still acts as a boundary (via is_heading in
    # units), but a Title-styled title page must not switch off text-based
    # detection for documents whose chapters are plain text.
    has_style_headings = has_chapter_heading_styles(doc)
    if has_style_headings:
        chapters = build_chapters(
            units, title, allow_all_caps=False, text_headings=False
        )
    else:
        explicit_count = sum(1 for text, _ in units if is_explicit_section(text))
        chapters = build_chapters(
            units,
            title,
            allow_all_caps=explicit_count < EXPLICIT_MARKER_THRESHOLD,
            text_headings=True,
        )

    if not chapters:
        chapters = [{
            'title': title,
            'order_num': 0,
            'content': raw.strip(),
            'word_count': len(raw.split()),
        }]

    return {
        'title': title,
        'author': author,
        'language': detect_language(raw),
        'cover_b64': None,
        'chapters': chapters,
    }
