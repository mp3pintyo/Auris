"""Tests for DOCX import.

Fixtures are built with python-docx in a temp directory so no binary
artifact is committed and the tests exercise a real round trip.
"""
import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from core.parser.docx_parser import DocxImportError, extract_units, parse


def _docx(build, name='book.docx'):
    """Build a .docx via a callback and return its path."""
    from docx import Document

    document = Document()
    build(document)
    path = Path(tempfile.mkdtemp()) / name
    document.save(path)
    return path


def _open(path):
    from docx import Document

    return Document(path)


class ExtractUnitsTests(unittest.TestCase):
    def test_heading_styles_are_flagged(self):
        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph('Body text.')

        units = extract_units(_open(_docx(build)))
        self.assertEqual(units[0], ('Chapter One', True))
        self.assertEqual(units[1], ('Body text.', False))

    def test_title_style_is_a_heading(self):
        def build(d):
            d.add_paragraph('The Book', style='Title')
            d.add_paragraph('Body text.')

        units = extract_units(_open(_docx(build)))
        self.assertEqual(units[0], ('The Book', True))

    def test_table_text_is_included_in_document_order(self):
        """doc.paragraphs silently omits table cells, which would drop part of
        the book with no warning."""
        def build(d):
            d.add_paragraph('Before the table.')
            table = d.add_table(rows=1, cols=2)
            table.cell(0, 0).text = 'Cell A'
            table.cell(0, 1).text = 'Cell B'
            d.add_paragraph('After the table.')

        texts = [text for text, _ in extract_units(_open(_docx(build)))]
        self.assertEqual(
            texts, ['Before the table.', 'Cell A', 'Cell B', 'After the table.']
        )

    def test_table_cells_are_never_headings(self):
        def build(d):
            table = d.add_table(rows=1, cols=1)
            table.cell(0, 0).text = 'CHAPTER ONE'

        units = extract_units(_open(_docx(build)))
        self.assertEqual(units, [('CHAPTER ONE', False)])

    def test_empty_paragraphs_are_kept_as_blank_units(self):
        """Blank lines separate paragraphs for downstream narration, so they
        must survive extraction rather than being dropped."""
        def build(d):
            d.add_paragraph('First.')
            d.add_paragraph('')
            d.add_paragraph('Second.')

        texts = [text for text, _ in extract_units(_open(_docx(build)))]
        self.assertEqual(texts, ['First.', '', 'Second.'])

    def test_docx_import_error_is_a_runtime_error(self):
        self.assertTrue(issubclass(DocxImportError, RuntimeError))

    def test_text_inside_a_content_control_is_included(self):
        """Word's built-in cover pages and template/form-generated documents
        wrap text in w:sdt/w:sdtContent blocks, which a body walk that only
        handles w:p and w:tbl silently drops."""
        from docx.oxml.ns import qn

        def build(d):
            d.add_paragraph('Before the control.')
            d.add_paragraph('After the control.')

        path = _docx(build)
        document = _open(path)
        body = document.element.body

        # python-docx has no high-level API for content controls, so build
        # the w:sdt/w:sdtContent XML directly.
        sdt = body.makeelement(qn('w:sdt'), {})
        sdt_content = body.makeelement(qn('w:sdtContent'), {})
        p = body.makeelement(qn('w:p'), {})
        r = body.makeelement(qn('w:r'), {})
        t = body.makeelement(qn('w:t'), {})
        t.text = 'Wrapped in a content control.'
        r.append(t)
        p.append(r)
        sdt_content.append(p)
        sdt.append(sdt_content)

        # Insert the sdt between the two existing paragraphs.
        body.insert(1, sdt)

        texts = [text for text, _ in extract_units(document)]
        self.assertEqual(
            texts,
            [
                'Before the control.',
                'Wrapped in a content control.',
                'After the control.',
            ],
        )

    def test_nested_table_text_is_included(self):
        def build(d):
            outer = d.add_table(rows=1, cols=1)
            outer.cell(0, 0).text = 'Outer cell text.'
            inner = outer.cell(0, 0).add_table(rows=1, cols=1)
            inner.cell(0, 0).text = 'Nested cell text.'

        texts = [text for text, _ in extract_units(_open(_docx(build)))]
        self.assertIn('Nested cell text.', texts)


LONG = (
    'This sentence exists to push the chapter body past the minimum length '
    'that the shared chapter builder requires before it will emit a chapter. '
) * 6


class DocxParseTests(unittest.TestCase):
    def test_splits_on_heading_styles(self):
        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)
            d.add_heading('Chapter Two', level=1)
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual([c['title'] for c in book['chapters']],
                         ['Chapter One', 'Chapter Two'])

    def test_style_headings_win_and_text_markers_are_ignored(self):
        """Either/or rule: with real styles present, a plain paragraph reading
        'Chapter Five' is prose, not a boundary."""
        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)
            d.add_paragraph('Chapter Five')
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        titles = [c['title'] for c in book['chapters']]
        self.assertEqual(titles, ['Chapter One'])
        self.assertIn('Chapter Five', book['chapters'][0]['content'])

    def test_short_first_chapter_under_a_real_heading_style_is_kept(self):
        """A section whose title came from a declared Heading style must not
        be subject to the TXT front-matter word-count guard: the boundary was
        declared by the author, not guessed from text."""
        def build(d):
            d.add_heading('The Arrival', level=1)
            d.add_paragraph('A short first chapter, forty words or so, ' * 4)
            d.add_heading('The Departure', level=1)
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        titles = [c['title'] for c in book['chapters']]
        self.assertIn('The Arrival', titles)
        self.assertIn('The Departure', titles)

    def test_title_style_alone_does_not_arm_style_detection(self):
        """A Title-styled title page is common in documents whose chapter
        headings are plain text. Title must still act as a boundary, but must
        not by itself switch off text-based detection."""
        def build(d):
            d.add_paragraph('My Great Novel', style='Title')
            d.add_paragraph('by Jane Austen')
            d.add_paragraph('Chapter One')
            d.add_paragraph(LONG)
            d.add_paragraph('Chapter Two')
            d.add_paragraph(LONG)
            d.add_paragraph('Chapter Three')
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual(len(book['chapters']), 3)

    def test_falls_back_to_text_markers_without_styles(self):
        def build(d):
            d.add_paragraph('Chapter One')
            d.add_paragraph(LONG)
            d.add_paragraph('Chapter Two')
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual([c['title'] for c in book['chapters']],
                         ['Chapter One', 'Chapter Two'])

    def test_core_properties_supply_title_and_author(self):
        def build(d):
            d.core_properties.title = 'The Fourteen Carat Car'
            d.core_properties.author = 'Jeno Rejto'
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual(book['title'], 'The Fourteen Carat Car')
        self.assertEqual(book['author'], 'Jeno Rejto')

    def test_metadata_falls_back_to_text_when_properties_empty(self):
        def build(d):
            # python-docx stamps author='python-docx' on every document it
            # creates, so clear it to exercise the text fallback.
            d.core_properties.author = ''
            d.add_paragraph('A Quiet Book')
            d.add_paragraph('by Jane Austen')
            d.add_paragraph('Chapter One')
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual(book['title'], 'A Quiet Book')
        self.assertEqual(book['author'], 'Jane Austen')

    def test_generator_author_is_not_treated_as_the_book_author(self):
        """python-docx writes author='python-docx' by default. A file produced
        by a script must not import as written by python-docx."""
        def build(d):
            d.add_paragraph('A Quiet Book')
            d.add_paragraph('by Jane Austen')
            d.add_paragraph('Chapter One')
            d.add_paragraph(LONG)

        # Deliberately left at the python-docx default.
        book = parse(_docx(build))
        self.assertNotEqual(book['author'], 'python-docx')
        self.assertEqual(book['author'], 'Jane Austen')

    def test_table_text_reaches_chapter_content(self):
        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)
            table = d.add_table(rows=1, cols=1)
            table.cell(0, 0).text = 'Smuggled in a table cell.'

        book = parse(_docx(build))
        joined = '\n'.join(c['content'] for c in book['chapters'])
        self.assertIn('Smuggled in a table cell.', joined)

    def test_returns_the_parser_contract(self):
        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual(
            set(book), {'title', 'author', 'language', 'cover_b64', 'chapters'}
        )
        self.assertIsNone(book['cover_b64'])
        self.assertEqual(book['language'], 'en')
        for chapter in book['chapters']:
            self.assertEqual(
                set(chapter), {'title', 'order_num', 'content', 'word_count'}
            )
            self.assertEqual(chapter['word_count'], len(chapter['content'].split()))

    def test_single_chapter_fallback_for_unstructured_text(self):
        def build(d):
            d.add_paragraph(LONG)

        book = parse(_docx(build))
        self.assertEqual(len(book['chapters']), 1)
        self.assertIn('shared chapter builder', book['chapters'][0]['content'])

    def test_matches_txt_chapter_count_for_the_same_content(self):
        """Pins the shared builder: a styleless DOCX and the same text as TXT
        must split identically."""
        from core.parser import txt_parser

        lines = ['Chapter One', LONG, 'Chapter Two', LONG]

        def build(d):
            for line in lines:
                d.add_paragraph(line)

        txt_path = Path(tempfile.mkdtemp()) / 'book.txt'
        txt_path.write_text('\n'.join(lines), encoding='utf-8')

        from_docx = parse(_docx(build))
        from_txt = txt_parser.parse(txt_path)

        self.assertEqual(
            [c['title'] for c in from_docx['chapters']],
            [c['title'] for c in from_txt['chapters']],
        )


class DocxErrorTests(unittest.TestCase):
    def test_legacy_doc_is_named_explicitly(self):
        path = Path(tempfile.mkdtemp()) / 'legacy.docx'
        path.write_bytes(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1' + b'\x00' * 64)

        with self.assertRaises(DocxImportError) as ctx:
            parse(path)
        self.assertIn('.docx', str(ctx.exception))
        self.assertIn('legacy', str(ctx.exception).lower())

    def test_corrupt_file_reports_unreadable(self):
        path = Path(tempfile.mkdtemp()) / 'bad.docx'
        path.write_bytes(b'this is not a zip archive')

        with self.assertRaises(DocxImportError) as ctx:
            parse(path)
        self.assertIn('corrupt', str(ctx.exception).lower())

    def test_document_without_text_is_rejected(self):
        def build(d):
            d.add_paragraph('')

        with self.assertRaises(DocxImportError) as ctx:
            parse(_docx(build))
        self.assertIn('no readable text', str(ctx.exception).lower())

    def test_whitespace_only_document_is_rejected(self):
        """An apparently-empty real document usually contains whitespace-only
        paragraphs, not a truly empty string."""
        def build(d):
            d.add_paragraph('   ')
            d.add_paragraph('\t')

        with self.assertRaises(DocxImportError) as ctx:
            parse(_docx(build))
        self.assertIn('no readable text', str(ctx.exception).lower())

    def test_xml_comment_in_body_does_not_crash_the_import(self):
        """lxml yields comments/processing instructions with a callable .tag
        rather than a string; child.tag.split(...) on one used to raise
        AttributeError."""
        from lxml import etree

        def build(d):
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)

        path = _docx(build)
        document = _open(path)
        body = document.element.body
        comment = etree.Comment('a stray XML comment')
        body.insert(0, comment)

        units = extract_units(document)
        self.assertIn('Chapter One', [text for text, _ in units])

    def test_missing_dependency_message(self):
        """When python-docx is absent the user gets an install hint, not an
        ImportError traceback."""
        import builtins

        real_import = builtins.__import__

        def fake_import(name, *args, **kwargs):
            if name == 'docx' or name.startswith('docx.'):
                raise ImportError('No module named docx')
            return real_import(name, *args, **kwargs)

        path = Path(tempfile.mkdtemp()) / 'x.docx'
        path.write_bytes(b'PK\x03\x04' + b'\x00' * 64)

        builtins.__import__ = fake_import
        try:
            with self.assertRaises(DocxImportError) as ctx:
                parse(path)
        finally:
            builtins.__import__ = real_import
        self.assertIn('pip install python-docx', str(ctx.exception))


class DocxImportRouteTests(unittest.TestCase):
    """End to end: POST a real .docx to /api/books/import and check the row."""

    def setUp(self):
        import app as app_module
        from core import database
        from core import settings as app_settings

        self.app_module = app_module
        self.database = database
        self.app_settings = app_settings
        self.tmp = tempfile.TemporaryDirectory()
        self._original = (
            database.DB_PATH,
            app_settings.SETTINGS_FILE,
            app_module.UPLOAD_DIR,
            app_module._startup_complete,
        )
        self.addCleanup(self._restore_globals)
        database.DB_PATH = str(Path(self.tmp.name) / 'reader.db')
        app_settings.SETTINGS_FILE = Path(self.tmp.name) / 'settings.json'
        app_module.UPLOAD_DIR = self.tmp.name
        app_module._startup_complete = True
        database.init_db()
        app_module.app.config['TESTING'] = True
        self.client = app_module.app.test_client()

    def _restore_globals(self):
        (
            self.database.DB_PATH,
            self.app_settings.SETTINGS_FILE,
            self.app_module.UPLOAD_DIR,
            self.app_module._startup_complete,
        ) = self._original

    def tearDown(self):
        self.tmp.cleanup()

    def test_docx_upload_is_imported_and_stored_as_docx(self):
        def build(d):
            d.core_properties.author = ''
            d.add_heading('Chapter One', level=1)
            d.add_paragraph(LONG)

        payload = io.BytesIO(Path(_docx(build)).read_bytes())

        # narration_mode=single avoids needing an LLM endpoint configured.
        with patch.object(self.app_module.threading, 'Thread'):
            response = self.client.post(
                '/api/books/import',
                data={'file': (payload, 'book.docx'), 'narration_mode': 'single'},
                content_type='multipart/form-data',
            )

        self.assertEqual(response.status_code, 200, response.get_data(as_text=True))
        with self.database.get_conn() as conn:
            book = conn.execute('SELECT * FROM books').fetchone()
        self.assertEqual(book['file_type'], 'docx')
        self.assertEqual(book['total_chapters'], 1)
        self.assertEqual(book['title'], 'Chapter One')


if __name__ == '__main__':
    unittest.main()
